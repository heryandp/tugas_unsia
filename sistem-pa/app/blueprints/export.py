import io
import os
from datetime import datetime

from flask import Blueprint, current_app, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from openpyxl import Workbook

from ..db import get_db


bp = Blueprint("export", __name__)


@bp.route("/export/<tipe>")
def export_data(tipe):
    conn = get_db()
    wb = Workbook()
    ws = wb.active
    ws.title = "Export Data"
    if tipe == "perkara":
        data = conn.execute("SELECT * FROM perkara").fetchall()
        headers = ["No Perkara", "Suami", "Istri", "Status", "Biaya", "Hasil Mediasi", "Tanggal Sidang"]
        ws.append(headers)
        for row in data:
            ws.append(
                [
                    row["no_perkara"],
                    row["nama_suami"],
                    row["nama_istri"],
                    row["status"],
                    row["biaya_daftar"],
                    row["hasil_mediasi"],
                    row["tanggal_sidang"],
                ]
            )
    elif tipe == "pegawai":
        data = conn.execute("SELECT * FROM users").fetchall()
        headers = ["Username", "Role", "Nama Lengkap", "NIP", "PIN 6 Digit"]
        ws.append(headers)
        for row in data:
            ws.append([row["username"], row["role"], row["nama_lengkap"], row["nip"], row["pin_6digit"]])
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(out, download_name=f"Export_{tipe}.xlsx", as_attachment=True)


@bp.route("/download_doc_data/<id>")
def download_doc_data(id):
    conn = get_db()
    data = conn.execute("SELECT * FROM perkara WHERE id=?", (id,)).fetchone()
    logo_res = conn.execute("SELECT nilai FROM pengaturan WHERE kunci='logo_pa_ikn'").fetchone()
    if not data:
        return "Data tidak ditemukan"

    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    header = doc.sections[0].header
    header_table = header.add_table(1, 2, Inches(6))
    header_table.autofit = False
    header_table.allow_autofit = False
    left_cell = header_table.cell(0, 0)
    left_paragraph = left_cell.paragraphs[0]

    if logo_res and logo_res["nilai"] and logo_res["nilai"].startswith("/uploads/"):
        logo_path = logo_res["nilai"].replace("/uploads/", "")
        logo_full_path = os.path.join(current_app.config["UPLOAD_FOLDER"], logo_path)
        if os.path.exists(logo_full_path):
            run = left_paragraph.add_run()
            run.add_picture(logo_full_path, width=Inches(0.8))

    right_cell = header_table.cell(0, 1)
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = right_paragraph.add_run("MAHKAMAH AGUNG REPUBLIK INDONESIA\n")
    run1.bold = True
    run1.font.size = Pt(12)
    run2 = right_paragraph.add_run("PENGADILAN AGAMA IBUKOTA NUSANTARA\n")
    run2.bold = True
    run2.font.size = Pt(14)
    run3 = right_paragraph.add_run("Jl. Sumbu Kebangsaan No. 1, IKN\n")
    run3.font.size = Pt(9)

    doc.add_paragraph("─" * 80)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("DATA PERKARA PERCERAIAN\n")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 77, 0)

    nomor = doc.add_paragraph()
    nomor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nomor_run = nomor.add_run(f"Nomor: {data['no_perkara']}\n\n")
    nomor_run.font.size = Pt(12)
    nomor_run.bold = True

    # Registration Details Table
    table_reg = doc.add_table(rows=0, cols=2)
    table_reg.style = 'Table Grid' # Or None for invisible borders, usually 'Normal' or None is better for this doc
    # Using 'None' (default) style for no borders usually, but let's try to just add rows
    # docx default table style usually has borders? No, usually grid has borders.
    # Let's manually set width if possible or just rely on autofit.
    
    def add_row(table, label, value):
        row = table.add_row()
        c1 = row.cells[0]
        c1.text = label
        c1.paragraphs[0].runs[0].bold = True
        c1.width = Inches(2.5)
        c2 = row.cells[1]
        c2.text = value
        
    add_row(table_reg, "Tanggal Pendaftaran:", f"{data['tanggal_daftar']}")
    # add_row(table_reg, "Biaya Pendaftaran:", f"Rp {data['biaya_daftar']}")
    # Petugas
    add_row(table_reg, "Petugas Pendaftar:", f"{data['nama_staff']}")
    
    doc.add_paragraph()

    doc.add_heading("DATA PIHAK BERPERKARA:", level=2)
    table_pihak = doc.add_table(rows=0, cols=2)
    add_row(table_pihak, "Nama Suami (Pemohon):", f"{data['nama_suami']}")
    add_row(table_pihak, "Nama Istri (Termohon):", f"{data['nama_istri']}")
    doc.add_paragraph()

    doc.add_heading("POKOK PERKARA:", level=2)
    doc.add_paragraph(data["keluhan"] if data["keluhan"] else "-")
    doc.add_paragraph()

    doc.add_heading("STATUS PERKARA:", level=2)
    table_status = doc.add_table(rows=0, cols=2)
    add_row(table_status, "Status Terakhir:", f"{data['status']}")
    add_row(table_status, "Hasil Mediasi:", f"{data['hasil_mediasi']}")
    if data["detail_mediasi"]:
         add_row(table_status, "Catatan Mediasi:", f"{data['detail_mediasi']}")
    if data["nama_mediator"]:
         add_row(table_status, "Mediator:", f"{data['nama_mediator']}")
    if data["tanggal_sidang"]:
         add_row(table_status, "Jadwal Sidang:", f"{data['tanggal_sidang']}")
    doc.add_paragraph()

    if data["no_akta_cerai"]:
        doc.add_heading("NOMOR AKTA CERAI:", level=2)
        doc.add_paragraph(data["no_akta_cerai"])

    doc.add_paragraph("\n\n")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Dicetak pada: {datetime.now().strftime('%d %B %Y, %H:%M WIB')}\n"
    )
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    footer_run2 = footer_para.add_run("Dokumen ini dicetak dari Sistem Informasi PA IKN")
    footer_run2.font.size = Pt(8)
    footer_run2.font.color.rgb = RGBColor(128, 128, 128)


    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"Data_Perkara_{data['no_perkara'].replace('/', '_')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@bp.route("/download_pdf_data/<id>")
def download_pdf_data(id):
    conn = get_db()
    data = conn.execute("SELECT * FROM perkara WHERE id=?", (id,)).fetchone()
    logo_res = conn.execute("SELECT nilai FROM pengaturan WHERE kunci='logo_pa_ikn'").fetchone()
    if not data:
        return "Data tidak ditemukan"
    pdf = FPDF("P", "mm", "A4")
    pdf.add_page()
    pdf.set_line_width(1)
    pdf.rect(5, 5, 200, 287)
    pdf.set_line_width(0.5)
    pdf.rect(8, 8, 194, 281)
    if logo_res and logo_res["nilai"] and logo_res["nilai"].startswith("/uploads/"):
        logo_path = logo_res["nilai"].replace("/uploads/", "")
        logo_full_path = os.path.join(current_app.config["UPLOAD_FOLDER"], logo_path)
        if os.path.exists(logo_full_path):
            pdf.image(logo_full_path, 15, 12, 25)
    pdf.set_xy(45, 12)
    pdf.set_font("Times", "B", 14)
    pdf.cell(0, 7, "MAHKAMAH AGUNG REPUBLIK INDONESIA", 0, 1, "L")
    pdf.set_x(45)
    pdf.set_font("Times", "B", 16)
    pdf.cell(0, 7, "PENGADILAN AGAMA IBUKOTA NUSANTARA", 0, 1, "L")
    pdf.set_x(45)
    pdf.set_font("Times", "", 10)
    pdf.cell(0, 5, "Jl. Sumbu Kebangsaan No. 1, IKN", 0, 1, "L")
    pdf.ln(5)
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(10)
    pdf.set_font("Times", "B", 18)
    pdf.set_text_color(0, 77, 0)
    pdf.cell(0, 10, "DATA PERKARA PERCERAIAN", 0, 1, "C")
    pdf.set_font("Courier", "B", 12)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, f"Nomor: {data['no_perkara']}", 0, 1, "C")
    pdf.ln(8)
    pdf.set_font("Times", "", 11)
    pdf.set_x(20)
    pdf.cell(0, 7, f"Tanggal Pendaftaran: {data['tanggal_daftar']}", 0, 1)
    pdf.set_x(20)
    pdf.cell(0, 7, f"Biaya Pendaftaran: Rp {data['biaya_daftar']}", 0, 1)
    pdf.set_x(20)
    pdf.cell(0, 7, f"Petugas Pendaftar: {data['nama_staff']}", 0, 1)
    pdf.ln(5)
    pdf.set_font("Times", "B", 12)
    pdf.set_x(20)
    pdf.cell(0, 8, "DATA PIHAK BERPERKARA:", 0, 1)
    pdf.set_font("Times", "", 11)
    pdf.set_x(25)
    pdf.cell(0, 7, f"Nama Suami (Pemohon): {data['nama_suami']}", 0, 1)
    pdf.set_x(25)
    pdf.cell(0, 7, f"Nama Istri (Termohon): {data['nama_istri']}", 0, 1)
    pdf.ln(5)
    pdf.set_font("Times", "B", 12)
    pdf.set_x(20)
    pdf.cell(0, 8, "POKOK PERKARA / KELUHAN:", 0, 1)
    pdf.set_font("Times", "", 11)
    pdf.set_x(25)
    keluhan = data["keluhan"] if data["keluhan"] else "-"
    pdf.multi_cell(170, 6, keluhan)
    pdf.ln(3)
    pdf.set_font("Times", "B", 12)
    pdf.set_x(20)
    pdf.cell(0, 8, "STATUS & PROSES PERKARA:", 0, 1)
    pdf.set_font("Times", "", 11)
    pdf.set_x(25)
    pdf.cell(0, 7, f"Status Terakhir: {data['status']}", 0, 1)
    pdf.set_x(25)
    pdf.cell(0, 7, f"Hasil Mediasi: {data['hasil_mediasi']}", 0, 1)
    if data["nama_mediator"]:
        pdf.set_x(25)
        pdf.cell(0, 7, f"Mediator: {data['nama_mediator']}", 0, 1)
    if data["detail_mediasi"]:
        pdf.set_x(25)
        pdf.multi_cell(170, 6, f"Catatan Mediasi: {data['detail_mediasi']}")
    if data["tanggal_sidang"]:
        pdf.set_x(25)
        pdf.cell(0, 7, f"Jadwal Sidang: {data['tanggal_sidang']}", 0, 1)
    pdf.ln(5)
    if data["no_akta_cerai"]:
        pdf.set_font("Times", "B", 12)
        pdf.set_x(20)
        pdf.cell(0, 8, "NOMOR AKTA CERAI:", 0, 1)
        pdf.set_font("Courier", "B", 11)
        pdf.set_x(25)
        pdf.cell(0, 7, data["no_akta_cerai"], 0, 1)
    pdf.ln(15)
    pdf.set_font("Times", "I", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, f"Dicetak pada: {datetime.now().strftime('%d %B %Y, %H:%M WIB')}", 0, 1, "C")
    pdf.set_font("Times", "I", 8)
    pdf.cell(0, 5, "Dokumen ini dicetak dari Sistem Informasi Pengadilan Agama IKN", 0, 1, "C")
    pdf.ln(5)
    pdf.set_font("Times", "I", 7)
    pdf.set_text_color(120, 120, 120)
    buffer = io.BytesIO()
    pdf_output = pdf.output(dest="S")
    if isinstance(pdf_output, (bytes, bytearray)):
        buffer.write(bytes(pdf_output))
    else:
        buffer.write(str(pdf_output).encode("latin-1"))
    buffer.seek(0)
    filename = f"Data_Perkara_{data['no_perkara'].replace('/', '_')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype="application/pdf")


@bp.route("/cetak_akta/<tipe>/<id>")
def cetak_akta(tipe, id):
    conn = get_db()
    data = conn.execute("SELECT * FROM perkara WHERE id=?", (id,)).fetchone()
    logo_res = conn.execute("SELECT nilai FROM pengaturan WHERE kunci='logo_pa_ikn'").fetchone()
    barcode_res = conn.execute("SELECT nilai FROM pengaturan WHERE kunci='barcode_ttd'").fetchone()
    if not data["no_akta_cerai"]:
        return "Nomor Akta Belum Terbit!"
    pdf = FPDF("P", "mm", "A4")
    pdf.add_page()
    pdf.set_line_width(1)
    pdf.rect(5, 5, 200, 287)
    pdf.set_line_width(0.5)
    pdf.rect(8, 8, 194, 281)
    if logo_res and logo_res["nilai"] and logo_res["nilai"].startswith("/uploads/"):
        logo_path = logo_res["nilai"].replace("/uploads/", "")
        logo_full_path = os.path.join(current_app.config["UPLOAD_FOLDER"], logo_path)
        if os.path.exists(logo_full_path):
            pdf.image(logo_full_path, 15, 12, 25)
    pdf.set_xy(45, 12)
    pdf.set_font("Times", "B", 14)
    pdf.cell(0, 7, "MAHKAMAH AGUNG REPUBLIK INDONESIA", 0, 1, "L")
    pdf.set_x(45)
    pdf.set_font("Times", "B", 16)
    pdf.cell(0, 7, "PENGADILAN AGAMA IBUKOTA NUSANTARA", 0, 1, "L")
    pdf.set_x(45)
    pdf.set_font("Times", "", 10)
    pdf.cell(0, 5, "Jl. Sumbu Kebangsaan No. 1, Kawasan Inti Pusat Pemerintahan, IKN", 0, 1, "L")
    pdf.ln(2)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(15)
    pdf.set_font("Times", "B", 24)
    pdf.cell(0, 10, "AKTA CERAI", 0, 1, "C")
    pdf.set_font("Courier", "B", 12)
    pdf.cell(0, 10, f"Nomor Akta: {data['no_akta_cerai']}", 0, 1, "C")
    pdf.ln(15)
    pdf.set_font("Times", "", 12)
    penerima = data["nama_suami"] if tipe == "suami" else data["nama_istri"]
    status_p = "MANTAN SUAMI" if tipe == "suami" else "MANTAN ISTRI"
    text_opening = (
        "Panitera Pengadilan Agama Ibukota Nusantara menerangkan bahwa "
        "pada hari ini telah diberikan Akta Cerai sebagai bukti sah perceraian antara:"
    )
    pdf.multi_cell(0, 8, text_opening, 0, "C")
    pdf.ln(5)
    pdf.set_font("Times", "B", 12)
    pdf.cell(0, 8, f"PIHAK SUAMI: {data['nama_suami']}", 0, 1, "C")
    pdf.cell(0, 8, "DAN", 0, 1, "C")
    pdf.cell(0, 8, f"PIHAK ISTRI: {data['nama_istri']}", 0, 1, "C")
    pdf.ln(5)
    pdf.set_font("Times", "", 12)
    text_closing = (
        f"Telah resmi bercerai berdasarkan Putusan Pengadilan Agama IKN Nomor {data['no_perkara']}. "
        f"Akta ini diberikan kepada {penerima} ({status_p})."
    )
    pdf.multi_cell(0, 8, text_closing, 0, "J")
    pdf.ln(30)
    # ... (inside cetak_akta) ...
    
    # Fetch Panitera Name
    panitera_user = conn.execute("SELECT nama_lengkap FROM users WHERE role='panitera'").fetchone()
    nama_panitera = panitera_user["nama_lengkap"] if panitera_user else "............................................."

    pdf.set_x(120)
    pdf.cell(0, 6, f"IKN, {datetime.now().strftime('%d %B %Y')}", 0, 1, "L")
    pdf.set_x(120)
    pdf.cell(0, 6, "Panitera,", 0, 1, "L")
    pdf.ln(25) # Space for signature
    pdf.set_x(120)
    pdf.set_font("Times", "B", 12)
    pdf.cell(0, 6, nama_panitera, 0, 1, "L")
    pdf.set_font("Times", "", 12) # Reset font
    
    # Check for barcode/digital signature BEFORE name
    if barcode_res and barcode_res["nilai"] and barcode_res["nilai"].startswith("/uploads/"):
        barcode_path = barcode_res["nilai"].replace("/uploads/", "")
        barcode_full_path = os.path.join(current_app.config["UPLOAD_FOLDER"], barcode_path)
        if os.path.exists(barcode_full_path):
            current_y = pdf.get_y()
            pdf.image(barcode_full_path, 120, current_y - 20, 30) # Overlay slightly above
        else:
            pdf.set_x(120)
            pdf.set_font("Courier", "B", 8)
            pdf.cell(0, 5, "[ BARCODE TANDA TANGAN DIGITAL ]", 0, 1, "L")
            pdf.ln(15)
    else:
        pdf.set_x(120)
        pdf.set_font("Courier", "B", 8)
        pdf.cell(0, 5, "[ BARCODE TANDA TANGAN DIGITAL ]", 0, 1, "L")
        pdf.ln(15)
    if logo_res and logo_res["nilai"] and logo_res["nilai"].startswith("/uploads/"):
        logo_path = logo_res["nilai"].replace("/uploads/", "")
        logo_full_path = os.path.join(current_app.config["UPLOAD_FOLDER"], logo_path)
        if os.path.exists(logo_full_path):
            pdf.image(logo_full_path, 120, pdf.get_y(), 20)
            pdf.ln(15)
    pdf.set_x(120)
    pdf.set_font("Times", "BU", 12)
    pdf.cell(0, 6, "( MATERAI & CAP BASAH )", 0, 1, "L")
    pdf.set_y(280)
    pdf.set_font("Times", "I", 7)
    pdf.set_text_color(120, 120, 120)
    buffer = io.BytesIO()
    pdf_output = pdf.output(dest="S")
    if isinstance(pdf_output, (bytes, bytearray)):
        buffer.write(bytes(pdf_output))
    else:
        buffer.write(str(pdf_output).encode("latin-1"))
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Akta_Cerai_{tipe}_{data['no_perkara'].replace('/', '_')}.pdf",
        mimetype="application/pdf",
    )
